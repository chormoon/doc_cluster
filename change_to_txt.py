import docx
import os
import re


def changeAllFile(base):
    for root, dirs, files in os.walk(base):
        for filenamedocx in files:
            #每个循环中docx文档和txt文档的命名
            # filenamedocx = (str(i)+'.docx')
            # 用正则匹配，去掉不需要的词
            newName = re.sub("-", "_", filenamedocx)
            newName2 = re.sub(" ", "_", newName)
            # 设置新文件名
            newFilename = filenamedocx.replace(filenamedocx, newName2)
            print('new file name:'+newFilename)
            os.chdir('E:/Work/BWD/DAY1/doc_cluster/data/')

            os.rename(filenamedocx, newFilename)

            suffix = newFilename.split(".")[0]

            filenametxt = (suffix+'.txt')
            print("txtname:"+filenametxt)
            #新建和打开txt文档
            f = open(filenametxt,'w')
            #打开docx的文档并读入名为file的变量中
            file_ = docx.Document(newFilename)
            #输入docx中的段落数，以检查是否空文档
            print('段落数:'+str(len(file_.paragraphs)))
            #将每个段落的内容都写进去txt里面
            for para in file_.paragraphs:
                f.write(para.text)
            f.close()
            print('修改成功:'+filenametxt)
        print('操作已经完成！')

def main():
    base = './data/'
    changeAllFile(base)

if __name__ == '__main__':
    main()